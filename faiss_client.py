import os
import numpy as np
import faiss
import json
from sentence_transformers import SentenceTransformer, util
from docx import Document

import nltk
nltk.download('punkt')
from nltk.tokenize import sent_tokenize


class FAISSClient:
    """
    A vector indexing client that processes .docx files, extracts semantically meaningful text chunks,
    embeds them using a SentenceTransformer model, and indexes them using FAISS for semantic search.
    """

    def __init__(self):
        """
        Initializes the FAISSClient with:
        - a SentenceTransformer model for embeddings,
        - an empty FAISS index (built later),
        - empty lists for vectors and their associated metadata.
        """
        self.model = SentenceTransformer("BAAI/bge-base-en-v1.5")
        self.index = None
        self.vectors = []
        self.metadata = []

    def semantic_chunking(self, text, similarity_threshold=0.65, max_words=400, min_words=20):
        """
        Splits a long text into semantically meaningful chunks based on sentence similarity.

        Parameters:
            text (str): The input paragraph or section.
            similarity_threshold (float): Cosine similarity threshold to split on.
            max_words (int): Maximum word count per chunk.
            min_words (int): Minimum word count to avoid very short chunks.

        Returns:
            List[str]: A list of semantically segmented text chunks.
        """
        sentences = sent_tokenize(text)
        semantic_chunks = []
        current_chunk = [sentences[0]]
        last_embedding = self.model.encode(sentences[0], convert_to_tensor=True)

        for sentence in sentences[1:]:
            emb = self.model.encode(sentence, convert_to_tensor=True)
            sim = util.cos_sim(last_embedding, emb).item()
            chunk_words = " ".join(current_chunk).split()

            if len(chunk_words) < min_words:
                current_chunk.append(sentence)
            elif len(chunk_words) > max_words:
                semantic_chunks.append(". ".join(current_chunk))
                current_chunk = [sentence]
            elif sim < similarity_threshold and len(current_chunk) > 1:
                semantic_chunks.append(". ".join(current_chunk))
                current_chunk = [sentence]
            else:
                current_chunk.append(sentence)

            last_embedding = emb

        if current_chunk:
            semantic_chunks.append(" ".join(current_chunk))

        return semantic_chunks

    def parse_docx(self, file_path):
        """
        Parses a .docx file into chunks using headings and paragraph structure.

        Parameters:
            file_path (str): Path to the Word (.docx) file.

        Returns:
            List[str]: All chunks added from the document.
        """
        doc = Document(file_path)
        main_title = doc.paragraphs[0].text.strip() if doc.paragraphs else "Untitled"

        current_subheading = None
        current_paragraph = []
        all_chunks = []

        for para in doc.paragraphs[1:]:
            text = para.text.strip()
            if not text:
                continue
            is_bold = all(run.bold for run in para.runs if run.text.strip())
            if is_bold:
                if current_subheading and current_paragraph:
                    all_chunks.extend(self._add_chunks(current_paragraph, current_subheading, main_title))
                current_subheading = text
                current_paragraph = []
            else:
                current_paragraph.append(text)

        if current_subheading and current_paragraph:
            all_chunks.extend(self._add_chunks(current_paragraph, current_subheading, main_title))

        return all_chunks

    def _add_chunks(self, paragraph, subheading, title):
        """
        Converts a paragraph and subheading into semantic chunks and embeds them.

        Parameters:
            paragraph (List[str]): Paragraph lines.
            subheading (str): The subheading title.
            title (str): Main document title.

        Returns:
            List[str]: The created chunk texts.
        """
        text = " ".join(paragraph)
        chunks = self.semantic_chunking(text)
        vectors = self.model.encode(chunks)
        for i, (vec, chunk_text) in enumerate(zip(vectors, chunks)):
            self.vectors.append(vec)
            self.metadata.append({
                "title": title,
                "subheading": subheading,
                "text": chunk_text,
                "text_place": i
            })
        return chunks

    def process_all(self, folder_path="./word_files"):
        """
        Processes all .docx files in a folder and builds a FAISS index.

        Parameters:
            folder_path (str): Path to the folder containing .docx files.
        """
        for filename in os.listdir(folder_path):
            if filename.endswith(".docx"):
                print(f"Processing {filename}...")
                self.parse_docx(os.path.join(folder_path, filename))
        self._build_index()

    def _build_index(self):
        """
        Builds the FAISS index from accumulated vectors.
        """
        print("🔧 Building FAISS index...")
        vector_array = np.array(self.vectors).astype("float32")
        self.index = faiss.IndexFlatL2(vector_array.shape[1])
        self.index.add(vector_array)
        print(f"✅ Indexed {len(self.vectors)} chunks.")

    def search(self, query, top_k=5):
        """
        Searches the FAISS index with a user query.

        Parameters:
            query (str): The user question.
            top_k (int): Number of top matching chunks to return.

        Returns:
            List[dict]: Metadata for the top_k most relevant chunks.
        """
        query_vector = self.model.encode([query]).astype("float32")
        distances, indices = self.index.search(query_vector, top_k)
        results = []
        for i in indices[0]:
            results.append(self.metadata[i])
        return results

    def save(self, index_path="faiss_index.index", metadata_path="faiss_metadata.json"):
        """
        Saves the FAISS index and metadata to disk.

        Parameters:
            index_path (str): Path to save the FAISS index.
            metadata_path (str): Path to save the metadata JSON file.
        """
        faiss.write_index(self.index, index_path)
        with open(metadata_path, "w", encoding="utf-8") as f:
            json.dump(self.metadata, f, indent=2, ensure_ascii=False)
        print(f"✅ Saved index to '{index_path}' and metadata to '{metadata_path}'")


# Optional: Execute indexing directly if run as a script
if __name__ == "__main__":
    client = FAISSClient()
    client.process_all("./data_files")
    client.save()
    print("FAISS index and metadata have been built and saved successfully.")
    
